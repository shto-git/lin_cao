"""Word document export with full formatting support.

Features: cover page, TOC, headers/footers, numbered headings,
          chapter references, quality report with color coding.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any


def export_full_document(
    outline: Any,
    drafts: list[Any],
    findings: list[Any] | None,
    output_path: str | Path,
    title: str = "林草规划文档",
    region: str = "",
    period: str = "",
    drafting_unit: str = "",
) -> str:
    """Export full planning document to Word (.docx).

    Args:
        outline: OutlineNode root with chapters/sections
        drafts: List of ChapterDraft objects
        findings: List of QualityFinding objects (optional)
        output_path: Path for .docx file
        title: Document title
        region: Planning region
        period: Planning period
        drafting_unit: Organization name

    Returns:
        Absolute path to saved file
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        raise ImportError(
            "Word 导出需要 python-docx。请运行: pip install python-docx"
        )

    doc = Document()

    # ── Page setup ──
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.5)

    # ── Style ──
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.first_line_indent = Pt(24)
    style.paragraph_format.line_spacing = 1.5

    # ── Cover Page ──
    for _ in range(6):
        doc.add_paragraph("")

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.font.size = Pt(22)
    run.bold = True
    run.font.name = "黑体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    doc.add_paragraph("")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("编制说明")
    sub_run.font.size = Pt(14)
    sub_run.font.name = "楷体"
    sub_run.element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")

    for _ in range(8):
        doc.add_paragraph("")

    # Info table
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = "Light Grid Accent 1"
    info_data = [("规划区域", region), ("规划期限", period), ("编制单位", drafting_unit)]
    for i, (label, value) in enumerate(info_data):
        info_table.cell(i, 0).text = label
        info_table.cell(i, 1).text = value or "（待填写）"

    # Page break
    doc.add_page_break()

    # ── Table of Contents placeholder ──
    toc_heading = doc.add_heading("目  录", level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("（请右键目录区域选择"更新域"以生成完整目录）")
    doc.add_page_break()

    # ── Chapters ──
    draft_map = {getattr(d, "outline_id", None): d for d in drafts}

    for chapter in outline.children:
        # Chapter heading
        ch_title = chapter.title
        ch_heading = doc.add_heading(ch_title, level=1)
        ch_heading.paragraph_format.space_before = Pt(18)

        for section in chapter.children:
            sec_title = section.title
            sec_heading = doc.add_heading(sec_title, level=2)
            sec_heading.paragraph_format.space_before = Pt(12)

            # Get draft content
            draft = draft_map.get(section.id)
            if draft:
                content_text = getattr(draft, "content", "") or draft.get("content", "")
                if content_text:
                    # Skip the "<!-- 模拟草稿 -->" prefix
                    if content_text.startswith("<!--"):
                        lines = content_text.split("
")
                        # Skip first comment line
                        content_text = "
".join(lines[1:]).strip()

                    for line in content_text.split("
"):
                        line = line.strip()
                        if not line or line.startswith("<!--"):
                            continue
                        elif line.startswith("**") and line.endswith("**"):
                            p = doc.add_paragraph()
                            run = p.add_run(line.strip("*"))
                            run.bold = True
                            run.font.size = Pt(11)
                        elif line.startswith(">"):
                            p = doc.add_paragraph(line.lstrip("> "))
                            p.paragraph_format.left_indent = Cm(0.5)
                        else:
                            p = doc.add_paragraph(line)
                            p.paragraph_format.first_line_indent = Pt(24)

                # References section
                evidence_ids = getattr(draft, "evidence_ids", None)
                if evidence_ids:
                    ref_para = doc.add_paragraph()
                    ref_run = ref_para.add_run("参考资料：")
                    ref_run.bold = True
                    ref_run.font.size = Pt(10)
                    for eid in evidence_ids:
                        ref_line = doc.add_paragraph(f"• {eid}")
                        ref_line.paragraph_format.left_indent = Cm(0.8)

            doc.add_page_break()

    # ── Quality Report ──
    if findings:
        doc.add_heading("质检报告", level=1)
        doc.add_paragraph(f"共发现 {len(findings)} 个问题")

        error_count = sum(1 for f in findings if getattr(f, "severity", "") == "error")
        warning_count = sum(1 for f in findings if getattr(f, "severity", "") == "warning")
        doc.add_paragraph(f"错误: {error_count}  |  警告: {warning_count}")

        for i, finding in enumerate(findings, 1):
            severity = getattr(finding, "severity", "info")
            code = getattr(finding, "code", "")
            message = getattr(finding, "message", "")
            location = getattr(finding, "location", "")
            suggestion = getattr(finding, "suggestion", "")

            p = doc.add_paragraph()
            severity_run = p.add_run(f"[{severity.upper()}] ")
            if severity == "error":
                severity_run.font.color.rgb = RGBColor(220, 38, 38)
            elif severity == "warning":
                severity_run.font.color.rgb = RGBColor(217, 119, 6)
            else:
                severity_run.font.color.rgb = RGBColor(59, 130, 246)
            severity_run.bold = True

            p.add_run(f"{code} @ {location}: {message}")

            if suggestion:
                sug_p = doc.add_paragraph()
                sug_p.paragraph_format.left_indent = Cm(0.5)
                sug_run = sug_p.add_run(f"建议: {suggestion}")
                sug_run.font.size = Pt(10)
                sug_run.font.color.rgb = RGBColor(5, 150, 105)

    # ── Save ──
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return str(output_path.absolute())
